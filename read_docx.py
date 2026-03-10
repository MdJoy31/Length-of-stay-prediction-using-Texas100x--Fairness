from docx import Document
import sys

doc = Document('Paper/Comparison_Table_Methodology_v2 (1).docx')

print("=" * 80)
print("PARAGRAPHS")
print("=" * 80)
for i, para in enumerate(doc.paragraphs):
    if para.text.strip():
        print(f"P{i}: {para.text}")

print("\n" + "=" * 80)
print("TABLES")
print("=" * 80)
for ti, table in enumerate(doc.tables):
    print(f"\n--- TABLE {ti} ---")
    for j, row in enumerate(table.rows):
        cells = [cell.text.strip() for cell in row.cells]
        sep = " | "
        print(f"  Row {j}: {sep.join(cells)}")
