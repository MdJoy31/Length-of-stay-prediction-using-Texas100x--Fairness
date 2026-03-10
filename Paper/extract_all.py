import docx
import fitz  # PyMuPDF
import os
import sys

sys.stdout.reconfigure(encoding='utf-8', errors='replace')

PAPER_DIR = r"D:\Research study\Research question ML\fairness_project_v2\fairness_project_v1\Paper"

# ==================== DOCX ====================
print("=" * 80)
print("DOCX: Comparison_Table_Methodology_v2 (1).docx")
print("=" * 80)
docx_path = os.path.join(PAPER_DIR, "Comparison_Table_Methodology_v2 (1).docx")
doc = docx.Document(docx_path)

print("\n--- Paragraphs ---")
for i, para in enumerate(doc.paragraphs):
    if para.text.strip():
        print(f"P{i}: {para.text}")

print("\n--- Tables ---")
for t_idx, table in enumerate(doc.tables):
    print(f"\n=== Table {t_idx+1} ({len(table.rows)} rows x {len(table.columns)} cols) ===")
    for r_idx, row in enumerate(table.rows):
        cells = [cell.text.strip().replace("\n", " | ") for cell in row.cells]
        sep = " ;; "
        print(f"  Row {r_idx}: {sep.join(cells)}")

# ==================== PDFs ====================
pdf_files = [
    "12913_2024_Article_11238.pdf",
    "3721201.3721373.pdf",
    "applsci-14-10523.pdf",
    "frai-06-1179226.pdf",
    "nihms-1911125.pdf",
    "Rachda_Naila_Mekhaldi_Journal_Information_Science_Engineering_2021.pdf",
    "ZJMA_11_2149318.pdf",
]

for pdf_name in pdf_files:
    pdf_path = os.path.join(PAPER_DIR, pdf_name)
    if not os.path.exists(pdf_path):
        print(f"\n{'=' * 80}")
        print(f"PDF NOT FOUND: {pdf_name}")
        continue

    print(f"\n{'=' * 80}")
    print(f"PDF: {pdf_name}")
    print("=" * 80)

    try:
        doc = fitz.open(pdf_path)
        num_pages = min(3, len(doc))
        for page_num in range(num_pages):
            page = doc[page_num]
            text = page.get_text()
            print(f"\n--- Page {page_num + 1} ---")
            print(text[:3000])  # First 3000 chars per page
            if len(text) > 3000:
                print(f"... [truncated, total {len(text)} chars on this page]")
        doc.close()
    except Exception as e:
        print(f"ERROR: {e}")

print("\n\nDONE.")
